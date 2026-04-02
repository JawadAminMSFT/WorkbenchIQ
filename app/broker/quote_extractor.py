"""
Quote Extractor Module

Extracts structured quote fields from carrier documents using LLM parsing.
"""

from __future__ import annotations

import asyncio
import io
import json
import re
from typing import Any, Dict, Optional, Tuple

from app.broker.cu_processor import BrokerCUProcessor
from app.broker.models import QuoteFields
from app.config import BrokerSettings, ContentUnderstandingSettings, OpenAISettings
from app.openai_client import chat_completion
from app.utils import setup_logging

logger = setup_logging()


def _extract_docx_text(file_content: bytes, file_name: str) -> str:
    """Extract text from a DOCX/DOC file.

    Handles modern OOXML (.docx) via python-docx and OLE2-wrapped/encrypted
    Office files via msoffcrypto-tool fallback.
    """
    # Check if the file is OLE2-wrapped (starts with D0 CF 11 E0)
    if file_content[:4] == b'\xd0\xcf\x11\xe0':
        logger.info(f"OLE2 container detected for {file_name}, attempting decryption")
        try:
            import msoffcrypto
            encrypted_file = msoffcrypto.OfficeFile(io.BytesIO(file_content))
            if encrypted_file.is_encrypted():
                encrypted_file.load_key(password="")
                decrypted = io.BytesIO()
                encrypted_file.decrypt(decrypted)
                decrypted.seek(0)
                file_content = decrypted.read()
                logger.info(f"Decrypted OLE2 file {file_name}: {len(file_content)} bytes")
        except ImportError:
            logger.warning("msoffcrypto-tool not installed — cannot decrypt OLE2 Office files")
        except Exception as e:
            logger.warning(f"OLE2 decryption failed for {file_name}: {e}")

    try:
        from docx import Document

        doc = Document(io.BytesIO(file_content))
        parts: list[str] = []

        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        if parts:
            return "\n".join(parts)
        logger.warning(f"python-docx extracted no text from {file_name}")
    except ImportError:
        logger.error("python-docx not installed — cannot parse DOCX files")
    except Exception as e:
        logger.warning(f"DOCX extraction failed for {file_name}: {e}")

    try:
        return file_content.decode("utf-8")
    except UnicodeDecodeError:
        return file_content.decode("latin-1")


def _extract_pdf_text(file_content: bytes, file_name: str) -> str:
    """Extract text from a PDF file using available PDF libraries.

    Tries PyPDF2 first; falls back to raw UTF-8 decode if unavailable.
    """
    # Try PyPDF2
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(io.BytesIO(file_content))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        if pages:
            return "\n\n".join(pages)
        logger.warning(f"PyPDF2 extracted no text from {file_name}, falling back to raw decode")
    except ImportError:
        logger.info("PyPDF2 not available, falling back to raw decode for PDF")
    except Exception as e:
        logger.warning(f"PyPDF2 failed for {file_name}: {e}, falling back to raw decode")

    # Fallback: raw decode
    try:
        return file_content.decode("utf-8")
    except UnicodeDecodeError:
        return file_content.decode("latin-1")


def _extract_excel_text(file_content: bytes, file_name: str) -> str:
    """Extract text from an Excel workbook using openpyxl.

    Iterates through all sheets and cells, producing a text representation
    that can be fed to the LLM extraction prompt.
    """
    try:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(file_content), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"=== Sheet: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                cell_values = [str(c) if c is not None else "" for c in row]
                if any(v.strip() for v in cell_values):
                    parts.append(" | ".join(cell_values))
        wb.close()
        if parts:
            return "\n".join(parts)
        logger.warning(f"openpyxl extracted no data from {file_name}")
    except ImportError:
        logger.error("openpyxl not installed — cannot parse Excel files")
    except Exception as e:
        logger.warning(f"Excel extraction failed for {file_name}: {e}")

    return f"[Unable to extract content from Excel file: {file_name}]"


# CU PascalCase → QuoteFields snake_case mapping
_CU_FIELD_MAP: Dict[str, str] = {
    "AnnualPremium": "annual_premium",
    "TotalInsuredValue": "total_insured_value",
    "BuildingLimit": "building_limit",
    "ContentsLimit": "contents_limit",
    "BusinessInterruptionLimit": "business_interruption_limit",
    "Deductibles": "deductible",
    "FloodSublimit": "flood_sublimit",
    "EarthquakeSublimit": "earthquake_sublimit",
    "NamedPerilsExclusions": "named_perils_exclusions",
    "SpecialConditions": "special_conditions",
    "PolicyPeriod": "policy_period",
    "EffectiveDate": "effective_date",
    "CarrierAMBestRating": "carrier_am_best_rating",
    "QuoteReferenceNumber": "quote_reference_number",
    "ExpirationDate": "expiry_date",
    "Underwriter": "underwriter",
    "QuoteDate": "quote_date",
    "CoverageForm": "coverage_form",
    "ValuationBasis": "valuation_basis",
    "CoinsuranceRequirement": "coinsurance_requirement",
    "PremiumBreakdown": "premium_breakdown",
}


class QuoteExtractor:
    """Extracts quote fields from carrier documents."""

    def __init__(
        self,
        settings: OpenAISettings,
        broker_settings: Optional[BrokerSettings] = None,
        cu_settings: Optional[ContentUnderstandingSettings] = None,
    ):
        """Initialize QuoteExtractor with OpenAI settings and optional CU settings.

        Args:
            settings: OpenAI configuration settings
            broker_settings: Optional broker CU analyzer settings
            cu_settings: Optional Content Understanding endpoint settings
        """
        self.settings = settings
        self.broker_settings = broker_settings
        self.cu_settings = cu_settings

    async def extract_quote(
        self, file_content: bytes, file_name: str, carrier_name: str
    ) -> Tuple[QuoteFields, Dict[str, float]]:
        """Extract quote fields from uploaded carrier document.

        Uses Azure Content Understanding as the primary extraction method when
        available, falling back to local text extraction + LLM parsing.

        Args:
            file_content: Raw bytes of the carrier document
            file_name: Original filename
            carrier_name: Name of the carrier providing the quote
            
        Returns:
            Tuple of (QuoteFields dataclass, confidence_scores dict)
            Confidence scores map field names to 0.0-1.0 values
        """
        # --- Primary path: Azure Content Understanding ---
        if self.cu_settings and self.broker_settings:
            try:
                processor = BrokerCUProcessor(self.cu_settings, self.broker_settings)
                if processor.is_available():
                    logger.info(
                        f"Using Content Understanding for quote extraction: {file_name}"
                    )
                    return await self._extract_via_cu(
                        file_content, file_name, carrier_name
                    )
            except Exception as exc:
                logger.warning(
                    f"CU extraction failed for {file_name}, falling back to LLM: {exc}",
                    exc_info=True,
                )

        # --- Fallback: local text extraction + LLM ---
        logger.info(f"Using LLM fallback for quote extraction: {file_name}")
        return await self._extract_via_llm(file_content, file_name, carrier_name)

    # ------------------------------------------------------------------ #
    #  CU-based extraction
    # ------------------------------------------------------------------ #

    async def _extract_via_cu(
        self, file_content: bytes, file_name: str, carrier_name: str
    ) -> Tuple[QuoteFields, Dict[str, float]]:
        """Extract quote fields via Azure Content Understanding.

        Args:
            file_content: Raw document bytes
            file_name: Original filename
            carrier_name: Name of the carrier

        Returns:
            Tuple of (QuoteFields, confidence dict keyed by snake_case field name)
        """
        processor = BrokerCUProcessor(self.cu_settings, self.broker_settings)
        analyzer_id = self.broker_settings.quote_analyzer

        result: Dict[str, Any] = await processor.analyze_document(
            file_content, file_name, "carrier_quote", analyzer_id
        )

        cu_fields = result.get("fields", {})
        cu_confidence = result.get("confidence", {})

        # Map CU PascalCase fields → snake_case for QuoteFields
        mapped: Dict[str, Any] = {}
        confidence: Dict[str, float] = {}
        for cu_key, snake_key in _CU_FIELD_MAP.items():
            if cu_key in cu_fields:
                value = cu_fields[cu_key]
                # Deductibles may come back as a list; join into a string
                if snake_key == "deductible" and isinstance(value, list):
                    value = "; ".join(str(v) for v in value)
                mapped[snake_key] = value
            if cu_key in cu_confidence:
                confidence[snake_key] = cu_confidence[cu_key]

        quote_fields = QuoteFields(
            annual_premium=mapped.get("annual_premium", ""),
            total_insured_value=mapped.get("total_insured_value", ""),
            building_limit=mapped.get("building_limit", ""),
            contents_limit=mapped.get("contents_limit", ""),
            business_interruption_limit=mapped.get("business_interruption_limit", ""),
            deductible=mapped.get("deductible", ""),
            flood_sublimit=mapped.get("flood_sublimit", ""),
            earthquake_sublimit=mapped.get("earthquake_sublimit", ""),
            named_perils_exclusions=mapped.get("named_perils_exclusions", []),
            special_conditions=mapped.get("special_conditions", []),
            policy_period=mapped.get("policy_period", ""),
            effective_date=mapped.get("effective_date"),
            carrier_am_best_rating=mapped.get("carrier_am_best_rating", ""),
            quote_reference_number=mapped.get("quote_reference_number", ""),
            expiry_date=mapped.get("expiry_date"),
            underwriter=mapped.get("underwriter", ""),
            quote_date=mapped.get("quote_date"),
            coverage_form=mapped.get("coverage_form", ""),
            valuation_basis=mapped.get("valuation_basis", ""),
            coinsurance_requirement=mapped.get("coinsurance_requirement", ""),
            premium_breakdown=mapped.get("premium_breakdown", []),
        )

        # Quality gate: if CU missed critical fields, fall back to LLM
        critical_missing = not mapped.get("annual_premium")
        if critical_missing:
            logger.warning(
                f"CU missed annual_premium for {carrier_name} — "
                "falling back to LLM for full extraction"
            )
            return await self._extract_via_llm(file_content, file_name, carrier_name)

        low = [f for f, s in confidence.items() if s < 0.60]
        if low:
            logger.warning(f"CU low-confidence fields for {carrier_name}: {low}")

        logger.info(
            f"CU extraction complete for {carrier_name} — "
            f"{len(mapped)} fields, {len(confidence)} confidence scores"
        )
        return quote_fields, confidence

    # ------------------------------------------------------------------ #
    #  LLM-based extraction (original approach)
    # ------------------------------------------------------------------ #

    async def _extract_via_llm(
        self, file_content: bytes, file_name: str, carrier_name: str
    ) -> Tuple[QuoteFields, Dict[str, float]]:
        """Extract quote fields using local text parsing + LLM chat completion."""
        file_ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""

        # Route to the appropriate text-extraction backend
        if file_ext == "pdf":
            text_content = await asyncio.to_thread(
                _extract_pdf_text, file_content, file_name
            )
        elif file_ext in ("xlsx", "xls"):
            text_content = await asyncio.to_thread(
                _extract_excel_text, file_content, file_name
            )
        elif file_ext in ("docx", "doc"):
            text_content = await asyncio.to_thread(
                _extract_docx_text, file_content, file_name
            )
        else:
            try:
                text_content = file_content.decode("utf-8")
            except UnicodeDecodeError:
                logger.warning(f"UTF-8 decode failed for {file_name}, trying latin-1")
                text_content = file_content.decode("latin-1")

        # Build extraction prompt
        prompt = self._build_extraction_prompt(text_content, carrier_name)

        # Call LLM in thread pool to avoid blocking
        result = await asyncio.to_thread(
            chat_completion,
            self.settings,
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert commercial insurance quote parser. Extract quote fields accurately and provide confidence scores.",
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0.0,
            max_tokens=4000,
        )

        raw_content = result.get("content", "")
        logger.debug(f"Received LLM response for quote extraction: {len(raw_content)} chars")

        # Parse JSON response
        parsed = self._parse_json_response(raw_content)

        # Map to QuoteFields dataclass
        fields_dict = parsed.get("fields", {})
        confidence = parsed.get("confidence", {})

        quote_fields = QuoteFields(
            annual_premium=fields_dict.get("annual_premium", ""),
            total_insured_value=fields_dict.get("total_insured_value", ""),
            building_limit=fields_dict.get("building_limit", ""),
            contents_limit=fields_dict.get("contents_limit", ""),
            business_interruption_limit=fields_dict.get("business_interruption_limit", ""),
            deductible=fields_dict.get("deductible", ""),
            flood_sublimit=fields_dict.get("flood_sublimit", ""),
            earthquake_sublimit=fields_dict.get("earthquake_sublimit", ""),
            named_perils_exclusions=fields_dict.get("named_perils_exclusions", []),
            special_conditions=fields_dict.get("special_conditions", []),
            policy_period=fields_dict.get("policy_period", ""),
            effective_date=fields_dict.get("effective_date"),
            carrier_am_best_rating=fields_dict.get("carrier_am_best_rating", ""),
            quote_reference_number=fields_dict.get("quote_reference_number", ""),
            expiry_date=fields_dict.get("expiry_date"),
            underwriter=fields_dict.get("underwriter", ""),
        )

        # Flag low-confidence fields
        low_confidence_fields = [
            field for field, score in confidence.items() if score < 0.60
        ]
        if low_confidence_fields:
            logger.warning(
                f"Low confidence fields for {carrier_name}: {low_confidence_fields}"
            )

        return quote_fields, confidence

    def _build_extraction_prompt(self, text_content: str, carrier_name: str) -> str:
        """Build the prompt for quote field extraction.
        
        Args:
            text_content: Document text
            carrier_name: Carrier providing the quote
            
        Returns:
            Formatted prompt string
        """
        return f"""Extract the following fields from this {carrier_name} insurance quote document.

**Document Content:**
{text_content[:15000]}

**Fields to Extract:**
1. annual_premium - Total annual premium amount
2. total_insured_value - Total insured value across all coverages
3. building_limit - Building coverage limit
4. contents_limit - Contents/BPP coverage limit
5. business_interruption_limit - Business interruption coverage limit
6. deductible - Deductible amount(s)
7. flood_sublimit - Flood coverage sublimit
8. earthquake_sublimit - Earthquake coverage sublimit
9. named_perils_exclusions - List of excluded perils or conditions
10. special_conditions - List of special policy conditions or endorsements
11. policy_period - Policy effective dates
12. effective_date - Policy effective / inception date (ISO 8601 format YYYY-MM-DD)
13. carrier_am_best_rating - AM Best financial strength rating
14. quote_reference_number - Quote or proposal reference number
15. expiry_date - Quote expiration date (ISO 8601 format YYYY-MM-DD)
16. underwriter - Name of underwriter

**Response Format:**
Return ONLY valid JSON with two keys:
- "fields": object with extracted field values (use empty string "" for missing text fields, empty array [] for list fields)
- "confidence": object with confidence scores 0.0-1.0 for each field

Example:
{{
  "fields": {{
    "annual_premium": "$125,000",
    "total_insured_value": "$5,000,000",
    "building_limit": "$3,000,000",
    "contents_limit": "$1,500,000",
    "business_interruption_limit": "$500,000",
    "deductible": "$5,000",
    "flood_sublimit": "$100,000",
    "earthquake_sublimit": "$250,000",
    "named_perils_exclusions": ["Flood outside sublimit", "Nuclear"],
    "special_conditions": ["Sprinkler warranty", "Alarm warranty"],
    "policy_period": "01/01/2025 - 01/01/2026",
    "effective_date": "2025-01-01",
    "carrier_am_best_rating": "A+",
    "quote_reference_number": "Q-2025-12345",
    "expiry_date": "2025-02-15",
    "underwriter": "John Smith"
  }},
  "confidence": {{
    "annual_premium": 0.95,
    "total_insured_value": 0.90,
    "building_limit": 0.88,
    "contents_limit": 0.85,
    "business_interruption_limit": 0.82,
    "deductible": 0.93,
    "flood_sublimit": 0.75,
    "earthquake_sublimit": 0.70,
    "named_perils_exclusions": 0.80,
    "special_conditions": 0.78,
    "policy_period": 0.95,
    "effective_date": 0.95,
    "carrier_am_best_rating": 0.65,
    "quote_reference_number": 0.98,
    "expiry_date": 0.92,
    "underwriter": 0.85
  }}
}}

Return ONLY the JSON, no markdown fences or additional text."""

    def _parse_json_response(self, raw_content: str) -> Dict:
        """Parse JSON from LLM response, handling code fences and comments.
        
        Args:
            raw_content: Raw LLM response text
            
        Returns:
            Parsed JSON dict (or error dict if parsing fails)
        """
        content_to_parse = raw_content.strip()

        # Strip markdown code fences
        if content_to_parse.startswith("```"):
            first_newline = content_to_parse.find("\n")
            if first_newline != -1:
                content_to_parse = content_to_parse[first_newline + 1 :]
            closing_idx = content_to_parse.rfind("```")
            if closing_idx != -1:
                content_to_parse = content_to_parse[:closing_idx].strip()

        try:
            parsed = json.loads(content_to_parse)
        except json.JSONDecodeError:
            # Strip JS-style comments and try again
            stripped = re.sub(r"//[^\n]*", "", content_to_parse)
            try:
                parsed = json.loads(stripped)
            except json.JSONDecodeError:
                logger.error(f"Failed to parse JSON response: {content_to_parse[:200]}")
                parsed = {
                    "fields": {},
                    "confidence": {},
                    "_error": "Failed to parse JSON response",
                }

        return parsed
