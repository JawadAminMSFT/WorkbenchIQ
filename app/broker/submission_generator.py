"""
Submission Generator Module

Extracts ACORD 125 and ACORD 140 fields from client submission documents.
"""

from __future__ import annotations

import asyncio
import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from app.broker.models import BrokerDocument
from app.config import BrokerSettings, ContentUnderstandingSettings, OpenAISettings
from app.openai_client import chat_completion
from app.utils import setup_logging

logger = setup_logging()

# ---------------------------------------------------------------------------
# CU PascalCase → existing snake_case field mappings
# ---------------------------------------------------------------------------

_ACORD_125_CU_FIELD_MAP: Dict[str, str] = {
    "ApplicantName": "applicant_name",
    "FEIN": "fein",
    "BusinessPhone": "business_phone",
    "MailingAddress": "mailing_address",
    "BusinessType": "business_type",
    "YearsInBusiness": "years_in_business",
    "SICCode": "sic_code",
    "NAICSCode": "naics_code",
    "AnnualGrossRevenue": "annual_gross_revenue",
    "NumberOfEmployees": "number_of_employees",
    "PriorCarrier": "prior_carrier",
    "PriorPremium": "prior_premium",
    "CurrentCoverageLimits": "current_coverage_limits",
    "CurrentDeductibles": "current_deductibles",
    "RequestedCoverageLimits": "requested_coverage_limits",
    "RequestedDeductibles": "requested_deductibles",
    "EffectiveDateRequested": "effective_date_requested",
    "LinesOfBusinessRequested": "lines_of_business_requested",
    "CoverageModifications": "coverage_modifications",
}

_ACORD_140_CU_FIELD_MAP: Dict[str, str] = {
    "PropertyLocations": "property_locations",
    "LossHistory": "loss_history",
    "PriorPolicyNumber": "prior_policy_number",
    "PriorExpirationDate": "prior_expiration_date",
    "TotalInsuredValue": "total_insured_value",
    "RequestedLimits": "requested_limits",
    "Deductibles": "deductibles",
    "SpecialConditions": "special_conditions",
    "MortgageeInfo": "mortgagee_info",
}


class SubmissionGenerator:
    """Extracts ACORD form fields from client submission documents."""

    def __init__(
        self,
        settings: OpenAISettings,
        broker_settings: Optional[BrokerSettings] = None,
        cu_settings: Optional[ContentUnderstandingSettings] = None,
    ):
        """Initialize SubmissionGenerator.

        Args:
            settings: OpenAI configuration settings (used for LLM fallback)
            broker_settings: Broker CU configuration (optional)
            cu_settings: Content Understanding endpoint settings (optional)
        """
        self.settings = settings
        self._broker_settings = broker_settings
        self._cu_settings = cu_settings

    async def extract_acord_fields(
        self, documents: List[BrokerDocument]
    ) -> Tuple[Dict[str, Any], Dict[str, Any], Dict[str, float], Dict[str, str]]:
        """Extract ACORD 125 and ACORD 140 fields from uploaded client documents.

        Uses Azure Content Understanding as the primary extraction path when
        configured and available, with automatic fallback to LLM-based extraction.

        Args:
            documents: List of BrokerDocument objects with client submission data

        Returns:
            Tuple of (acord_125_fields, acord_140_fields, confidence_scores, field_sources)
            - acord_125_fields: Dict with ACORD 125 General Application fields
            - acord_140_fields: Dict with ACORD 140 Property Section fields
            - confidence_scores: Dict mapping field names to confidence 0.0-1.0
            - field_sources: Dict mapping field names to source document filenames
        """
        logger.info(f"Extracting ACORD fields from {len(documents)} documents")

        # --- Primary path: Azure Content Understanding ---
        if self._broker_settings and self._cu_settings:
            try:
                cu_result = await self._extract_via_cu(documents)
                if cu_result:
                    acord_125 = cu_result["acord_125"]
                    acord_140 = cu_result["acord_140"]
                    confidence = cu_result["confidence"]
                    field_sources = cu_result["field_sources"]

                    low_confidence = [
                        f for f, s in confidence.items() if s < 0.60
                    ]
                    if low_confidence:
                        logger.warning(f"Low confidence ACORD fields (CU): {low_confidence}")
                    return acord_125, acord_140, confidence, field_sources
            except Exception:
                logger.exception(
                    "CU extraction failed — falling back to LLM extraction"
                )

        # --- Fallback: LLM-based extraction ---
        return await self._extract_via_llm(documents)

    # ------------------------------------------------------------------
    # CU-based extraction
    # ------------------------------------------------------------------

    async def _extract_via_cu(
        self, documents: List[BrokerDocument]
    ) -> Optional[Dict[str, Any]]:
        """Extract ACORD fields using Azure Content Understanding.

        Each document is analyzed individually via CU, then results are
        merged and mapped to the canonical snake_case field format.

        Returns:
            Dict with keys acord_125, acord_140, confidence, field_sources,
            or ``None`` if CU is unavailable.
        """
        from app.broker.cu_processor import BrokerCUProcessor
        from app.broker.cu_router import BrokerDocRouter

        processor = BrokerCUProcessor(
            cu_settings=self._cu_settings,
            broker_settings=self._broker_settings,
        )

        if not processor.is_available():
            logger.info("CU not available — skipping CU extraction path")
            return None

        logger.info("Using Content Understanding for ACORD field extraction")

        router = BrokerDocRouter()

        # Aggregate raw CU fields, confidences, and source tracking
        all_fields: Dict[str, Any] = {}
        all_confidence: Dict[str, float] = {}
        all_sources: Dict[str, str] = {}

        for doc in documents:
            # Load file content: try blob storage first, then local file path
            file_content = None
            if doc.blob_url:
                try:
                    from app.broker.storage import _is_blob_mode, _get_container_client, _BLOB_PREFIX
                    if _is_blob_mode() and not Path(doc.blob_url).exists():
                        # blob_url is a blob key — download from blob storage
                        from azure.core.exceptions import ResourceNotFoundError
                        cc = _get_container_client()
                        try:
                            file_content = cc.get_blob_client(doc.blob_url).download_blob().readall()
                        except ResourceNotFoundError:
                            pass
                except Exception:
                    pass

                if file_content is None:
                    fp = Path(doc.blob_url)
                    if fp.exists():
                        file_content = fp.read_bytes()

            if file_content is None:
                logger.warning(
                    "Skipping document '%s' — file not found at '%s'",
                    doc.file_name,
                    doc.blob_url,
                )
                continue

            # Auto-detect type from filename when doc type is "other"
            doc_type = doc.document_type
            if doc_type in ("other", "", None):
                detected = router._detect_type_from_filename(doc.file_name)
                if detected != "other":
                    logger.info(
                        "Auto-detected document '%s' as '%s' (was '%s')",
                        doc.file_name, detected, doc_type,
                    )
                    doc_type = detected

            analyzer_id = router.get_analyzer_id(
                doc_type, self._broker_settings
            )

            try:
                result = await processor.analyze_document(
                    file_content=file_content,
                    file_name=doc.file_name,
                    document_type=doc_type,
                    analyzer_id=analyzer_id,
                )
            except Exception:
                logger.warning(
                    "CU analysis failed for document '%s', skipping",
                    doc.file_name,
                    exc_info=True,
                )
                continue

            fields = result.get("fields", {})
            confidence = result.get("confidence", {})

            for cu_key, value in fields.items():
                if value is None or value == "":
                    continue
                # Only overwrite if new value has higher confidence
                existing_conf = all_confidence.get(cu_key, -1.0)
                new_conf = confidence.get(cu_key, 0.0)
                if new_conf >= existing_conf:
                    all_fields[cu_key] = value
                    all_confidence[cu_key] = new_conf
                    all_sources[cu_key] = doc.file_name

        if not all_fields:
            logger.warning("CU returned no fields — will fall back to LLM")
            return None

        # Map CU PascalCase fields to snake_case ACORD 125 / 140 dicts
        acord_125: Dict[str, Any] = {}
        acord_140: Dict[str, Any] = {}
        mapped_confidence: Dict[str, float] = {}
        mapped_sources: Dict[str, str] = {}

        for cu_key, value in all_fields.items():
            conf = all_confidence.get(cu_key, 0.0)
            source = all_sources.get(cu_key, "")

            if cu_key in _ACORD_125_CU_FIELD_MAP:
                snake = _ACORD_125_CU_FIELD_MAP[cu_key]
                acord_125[snake] = value
                mapped_confidence[snake] = conf
                if source:
                    mapped_sources[snake] = source
            elif cu_key in _ACORD_140_CU_FIELD_MAP:
                snake = _ACORD_140_CU_FIELD_MAP[cu_key]
                acord_140[snake] = value
                mapped_confidence[snake] = conf
                if source:
                    mapped_sources[snake] = source

        logger.info(
            "CU extraction complete: %d ACORD-125 fields, %d ACORD-140 fields",
            len(acord_125),
            len(acord_140),
        )

        # Quality gate: CU on plain text files often misses scalar fields
        # (only extracts arrays like PropertyLocations / LossHistory).
        # If fewer than 3 ACORD-125 fields were extracted, fall back to LLM
        # which handles unstructured text much better.
        if len(acord_125) < 3:
            logger.warning(
                "CU extracted only %d ACORD-125 fields (minimum 3 required) "
                "— falling back to LLM for better scalar field extraction",
                len(acord_125),
            )
            return None

        return {
            "acord_125": acord_125,
            "acord_140": acord_140,
            "confidence": mapped_confidence,
            "field_sources": mapped_sources,
        }

    # ------------------------------------------------------------------
    # LLM-based extraction (fallback)
    # ------------------------------------------------------------------

    async def _extract_via_llm(
        self, documents: List[BrokerDocument]
    ) -> Tuple[Dict[str, Any], Dict[str, Any], Dict[str, float], Dict[str, str]]:
        """Extract ACORD fields using LLM chat completion (fallback path).

        Returns:
            Tuple of (acord_125, acord_140, confidence, field_sources)
        """
        logger.info("Using LLM fallback for ACORD field extraction")

        # Aggregate document content
        document_context = self._build_document_context(documents)

        # Build extraction prompt
        prompt = self._build_extraction_prompt(document_context)

        # Call LLM in thread pool
        result = await asyncio.to_thread(
            chat_completion,
            self.settings,
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert commercial insurance submission analyst. Extract ACORD form fields accurately from client documents.",
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0.0,
            max_tokens=8000,
        )

        raw_content = result.get("content", "")
        logger.debug(f"Received LLM response for ACORD extraction: {len(raw_content)} chars")

        # Parse JSON response
        parsed = self._parse_json_response(raw_content)

        # Extract structured data
        acord_125 = parsed.get("acord_125", {})
        acord_140 = parsed.get("acord_140", {})
        confidence = parsed.get("confidence", {})
        field_sources = parsed.get("field_sources", {})

        # Log low-confidence fields
        low_confidence = [field for field, score in confidence.items() if score < 0.60]
        if low_confidence:
            logger.warning(f"Low confidence ACORD fields: {low_confidence}")

        return acord_125, acord_140, confidence, field_sources

    def _build_document_context(self, documents: List[BrokerDocument]) -> str:
        """Build aggregated document context for extraction.

        Reads actual file content from storage (local or blob) so the LLM
        has real data to extract ACORD fields from.
        """
        from pathlib import Path
        import io as _io

        context_parts = []

        for doc in documents:
            doc_info = f"\n=== Document: {doc.file_name} (Type: {doc.document_type}) ===\n"

            # Load raw bytes from storage (blob or local)
            raw = None
            if doc.blob_url:
                try:
                    from app.broker.storage import _is_blob_mode, _get_container_client
                    if _is_blob_mode() and not Path(doc.blob_url).exists():
                        from azure.core.exceptions import ResourceNotFoundError
                        cc = _get_container_client()
                        try:
                            raw = cc.get_blob_client(doc.blob_url).download_blob().readall()
                        except ResourceNotFoundError:
                            pass
                except Exception:
                    pass
                if raw is None:
                    fp = Path(doc.blob_url)
                    if fp.exists():
                        raw = fp.read_bytes()

            file_content = None
            if raw:
                try:
                    file_ext = Path(doc.file_name).suffix.lower()
                    if file_ext == ".pdf":
                        try:
                            from PyPDF2 import PdfReader
                            reader = PdfReader(_io.BytesIO(raw))
                            pages = [p.extract_text() for p in reader.pages if p.extract_text()]
                            file_content = "\n\n".join(pages) if pages else None
                        except ImportError:
                            file_content = raw.decode("utf-8", errors="replace")
                    elif file_ext == ".docx":
                        try:
                            from docx import Document
                            ddoc = Document(_io.BytesIO(raw))
                            parts = [p.text for p in ddoc.paragraphs if p.text.strip()]
                            for tbl in ddoc.tables:
                                for row in tbl.rows:
                                    cells = [c.text.strip() for c in row.cells]
                                    if any(cells):
                                        parts.append(" | ".join(cells))
                            file_content = "\n".join(parts) if parts else None
                        except ImportError:
                            file_content = raw.decode("utf-8", errors="replace")
                    else:
                        file_content = raw.decode("utf-8", errors="replace")
                    if file_content:
                        logger.info(f"Read {len(file_content)} chars from {doc.file_name}")
                except Exception as exc:
                    logger.warning(f"Could not read file {doc.file_name}: {exc}")

            if file_content:
                # Truncate to avoid exceeding token limits per document
                doc_info += file_content[:8000]
                doc_info += "\n"
            elif doc.extracted_fields:
                doc_info += "Extracted Fields:\n"
                doc_info += json.dumps(doc.extracted_fields, indent=2)[:5000]
                doc_info += "\n"
            else:
                doc_info += "(No content available for this document)\n"

            context_parts.append(doc_info)

        return "\n".join(context_parts)

    def _build_extraction_prompt(self, document_context: str) -> str:
        """Build the ACORD field extraction prompt.

        Args:
            document_context: Aggregated document content

        Returns:
            Formatted prompt string
        """
        return f"""Extract ACORD 125 (General Application) and ACORD 140 (Property Section) fields from these client submission documents.

**Document Context:**
{document_context[:25000]}

**ACORD 125 Fields to Extract (18 fields per spec):**
1. applicant_name - Named insured / business name
2. fein - Federal Employer Identification Number
3. business_phone - Primary contact phone
4. mailing_address - Full mailing address
5. business_type - Type of business/entity (Corporation, LLC, etc.)
6. years_in_business - Years the business has been operating
7. sic_code - Standard Industry Classification code
8. annual_gross_revenue - Annual gross revenue
9. number_of_employees - Full-time + part-time headcount
10. prior_carrier - Current/prior insurance carrier name
11. prior_premium - Current/prior annual premium
12. current_coverage_limits - Current policy coverage limits (e.g., per-occurrence, aggregate)
13. current_deductibles - Current policy deductible amounts
14. requested_coverage_limits - Requested coverage limits for the new policy
15. requested_deductibles - Requested deductible amounts for the new policy
16. coverage_modifications - Any requested coverage changes or endorsements
17. effective_date_requested - Requested policy effective date (ISO format YYYY-MM-DD)
18. lines_of_business_requested - List of coverage types (Property, GL, Umbrella, etc.)

**ACORD 140 Fields to Extract (15+ fields per spec):**
1. property_locations - Array of property locations with:
   - address: Full street address
   - occupancy: Building occupancy/use
   - construction_type: Construction class (frame, masonry, fire-resistive, etc.)
   - year_built: Year building was constructed
   - square_footage: Total building square footage
   - building_value: Building replacement value
   - contents_value: Business personal property value
   - bi_value: Business income/interruption value
   - protection_class: Fire protection class (1-10)
2. loss_history - Property loss history as an array of objects, each with:
   - date: Date of loss (ISO format YYYY-MM-DD)
   - cause: Cause of loss (e.g., "Water Damage", "Fire", "Wind/Hail")
   - amount: Total amount paid or reserved (dollar string)
   - description: Brief description of the loss event
   - status: "Open" or "Closed"
   - location: Property address where loss occurred (if available)
3. prior_policy_number - Expiring policy number
4. prior_expiration_date - Current policy expiry date
5. total_insured_value - Sum of all property values
6. requested_limits - Requested coverage limits (building, contents, BI)
7. deductibles - Requested deductible amounts
8. special_conditions - Any special conditions or requirements
9. mortgagee_info - Mortgagee/additional insured information

**Response Format:**
Return ONLY valid JSON with four keys:
{{
  "acord_125": {{
    "applicant_name": "Business Name Inc.",
    "fein": "12-3456789",
    "business_phone": "(555) 123-4567",
    "mailing_address": "123 Main St, City, ST 12345",
    "business_type": "Corporation",
    "years_in_business": 15,
    "sic_code": "5411",
    "annual_gross_revenue": "$12,500,000",
    "number_of_employees": 85,
    "prior_carrier": "ABC Insurance Co.",
    "prior_premium": "$85,000",
    "current_coverage_limits": "$1,000,000 per occurrence / $2,000,000 aggregate",
    "current_deductibles": "$5,000 property / $10,000 wind-hail",
    "requested_coverage_limits": "$2,000,000 per occurrence / $4,000,000 aggregate",
    "requested_deductibles": "$5,000 property / $25,000 wind-hail",
    "coverage_modifications": "Add equipment breakdown coverage; increase BI limit",
    "effective_date_requested": "2025-01-01",
    "lines_of_business_requested": ["Commercial Property", "General Liability", "Business Auto", "Umbrella"]
  }},
  "acord_140": {{
    "property_locations": [
      {{
        "address": "123 Main St, City, ST 12345",
        "occupancy": "Office",
        "construction_type": "Masonry Noncombustible",
        "year_built": 1995,
        "square_footage": 10000,
        "building_value": "$2,500,000",
        "contents_value": "$500,000",
        "bi_value": "$750,000",
        "protection_class": "3"
      }}
    ],
    "loss_history": [
      {{"date": "2023-05-15", "cause": "Water Damage", "amount": "$12,500", "description": "Roof leak during storm", "status": "Closed", "location": "123 Main St"}}
    ],
    "prior_policy_number": "POL-123456",
    "prior_expiration_date": "2024-12-31",
    "total_insured_value": "$8,250,000",
    "requested_limits": {{
      "building": "$5,500,000",
      "contents": "$1,500,000",
      "business_interruption": "$1,250,000"
    }},
    "deductibles": {{
      "property": "$5,000",
      "wind_hail": "$25,000"
    }},
    "special_conditions": [
      "Sprinkler system required",
      "24/7 security monitoring in place"
    ],
    "mortgagee_info": [
      {{"name": "First National Bank", "loan_number": "LOAN-789456", "address": "789 Bank St, City, ST 12345"}}
    ]
  }},
  "confidence": {{
    "applicant_name": 0.98,
    "fein": 0.95,
    "business_phone": 0.93,
    "mailing_address": 0.95,
    "business_type": 0.90,
    "years_in_business": 0.85,
    "sic_code": 0.88,
    "annual_gross_revenue": 0.87,
    "number_of_employees": 0.82,
    "prior_carrier": 0.88,
    "prior_premium": 0.82,
    "current_coverage_limits": 0.80,
    "current_deductibles": 0.78,
    "requested_coverage_limits": 0.82,
    "requested_deductibles": 0.80,
    "coverage_modifications": 0.70,
    "effective_date_requested": 0.92,
    "lines_of_business_requested": 0.90,
    "property_locations": 0.88,
    "loss_history": 0.78,
    "prior_policy_number": 0.85,
    "prior_expiration_date": 0.90,
    "total_insured_value": 0.92,
    "requested_limits": 0.85,
    "deductibles": 0.80,
    "special_conditions": 0.70,
    "mortgagee_info": 0.75
  }},
  "field_sources": {{
    "applicant_name": "application-form.pdf",
    "fein": "application-form.pdf",
    "property_locations": "schedule-of-values.xlsx",
    "loss_history": "loss-runs-2024.pdf",
    "prior_carrier": "prior-declaration.pdf",
    "current_coverage_limits": "prior-declaration.pdf",
    "current_deductibles": "prior-declaration.pdf"
  }}
}}

**field_sources**: Map each field name to the filename of the source document it was extracted from.
Use the document filenames exactly as provided in the Document Context headers above (e.g., "ridgeview-sov.txt").
Only include fields where the source document is identifiable.

Use empty string "" for missing text fields, empty array [] for missing lists, and null for missing objects.
Return ONLY the JSON, no markdown fences or additional text."""

    def _parse_json_response(self, raw_content: str) -> Dict[str, Any]:
        """Parse JSON from LLM response, handling code fences and comments.

        Args:
            raw_content: Raw LLM response text

        Returns:
            Parsed JSON dict (or error dict if parsing fails)
        """
        content_to_parse = raw_content.strip()

        # Strip markdown code fences
        if content_to_parse.startswith("```"):
            first_newline = content_to_parse.find("\n")
            if first_newline != -1:
                content_to_parse = content_to_parse[first_newline + 1 :]
            closing_idx = content_to_parse.rfind("```")
            if closing_idx != -1:
                content_to_parse = content_to_parse[:closing_idx].strip()

        # Strip JS-style comments
        content_to_parse = re.sub(r"//[^\n]*", "", content_to_parse)

        try:
            parsed = json.loads(content_to_parse)

            # Validate structure
            if "acord_125" not in parsed:
                parsed["acord_125"] = {}
            if "acord_140" not in parsed:
                parsed["acord_140"] = {}
            if "confidence" not in parsed:
                parsed["confidence"] = {}
            if "field_sources" not in parsed:
                parsed["field_sources"] = {}

            return parsed

        except json.JSONDecodeError as exc:
            logger.error(f"Failed to parse ACORD JSON: {exc}")
            return {
                "acord_125": {},
                "acord_140": {},
                "confidence": {},
                "_error": f"Failed to parse JSON response: {str(exc)}",
                "_raw": raw_content[:500],
            }
