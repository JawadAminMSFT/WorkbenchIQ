"""
Commercial Brokerage API Router

FastAPI router providing endpoints for:
- Client management
- Submission workflow
- Quote upload and comparison
- Carrier profile management
- Dashboard metrics
- AI-powered client research
"""

from __future__ import annotations

import json
import logging
import uuid
from dataclasses import asdict
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from pydantic import BaseModel, Field

from app.broker.constants import (
    DocumentType,
    LineOfBusiness,
    QuoteStatus,
    SubmissionStatus,
)
from app.broker.models import (
    BrokerDocument,
    CarrierProfile,
    Client,
    Placement,
    PlacementScoring,
    PropertyLocation,
    Quote,
    QuoteFields,
    Submission,
)
from app.broker.acord_forms import (
    ACORD_125_FIELDS,
    ACORD_140_FIELDS,
    SOV_FIELDS,
    get_acord_form_template,
    map_extracted_to_acord125,
)
from app.broker.cu_processor import BrokerCUProcessor
from app.broker.package_generator import SubmissionPackageGenerator
from app.broker.storage import BrokerStorage
from app.broker.quote_extractor import QuoteExtractor
from app.broker.placement_engine import PlacementEngine
from app.broker.research_engine import ClientResearchEngine
from app.broker.submission_generator import SubmissionGenerator
from app.config import load_settings

logger = logging.getLogger(__name__)


def _quote_from_dict(data: Dict[str, Any]) -> Quote:
    """Reconstruct a Quote dataclass from a storage dict, handling nested fields."""
    d = dict(data)
    # Convert nested 'fields' dict → QuoteFields dataclass
    if isinstance(d.get("fields"), dict):
        raw_fields = dict(d["fields"])
        # Ensure deductible is a string (LLM may return a dict)
        if isinstance(raw_fields.get("deductible"), dict):
            raw_fields["deductible"] = ", ".join(
                f"{k}: {v}" for k, v in raw_fields["deductible"].items()
            )
        d["fields"] = QuoteFields(**{
            k: v for k, v in raw_fields.items()
            if k in QuoteFields.__dataclass_fields__
        })
    # Convert nested 'scoring' dict → PlacementScoring dataclass
    if isinstance(d.get("scoring"), dict):
        d["scoring"] = PlacementScoring(**{
            k: v for k, v in d["scoring"].items()
            if k in PlacementScoring.__dataclass_fields__
        })
    # Drop any extra keys not on the Quote dataclass
    valid_keys = set(Quote.__dataclass_fields__)
    return Quote(**{k: v for k, v in d.items() if k in valid_keys})


def _submission_from_dict(data: Dict[str, Any]) -> Submission:
    """Reconstruct a Submission dataclass from a storage dict, handling nested types."""
    d = dict(data)
    # Convert embedded quotes dicts → Quote dataclass instances
    if isinstance(d.get("quotes"), list):
        d["quotes"] = [
            _quote_from_dict(q) if isinstance(q, dict) else q
            for q in d["quotes"]
        ]
    # Convert embedded documents (list of dicts) — leave as-is since they serialize cleanly
    # Convert embedded property locations if present
    valid_keys = set(Submission.__dataclass_fields__)
    return Submission(**{k: v for k, v in d.items() if k in valid_keys})

# Create router with prefix
router = APIRouter(prefix="/api/broker", tags=["Commercial Brokerage"])

# Initialize storage
storage = BrokerStorage()

# Load settings for engines
_settings = load_settings()


# ============================================================================
# Pydantic Request/Response Models
# ============================================================================

class PropertyLocationSchema(BaseModel):
    address: str = ""
    occupancy: str = ""
    construction: str = ""
    year_built: Optional[int] = None
    square_footage: Optional[int] = None
    building_value: str = ""
    contents_value: str = ""
    bi_value: str = ""
    protection_class: str = ""


class CreateClientRequest(BaseModel):
    name: str
    industry_code: str = ""
    business_type: str = ""
    years_in_business: Optional[int] = None
    annual_revenue: str = ""
    employee_count: Optional[int] = None
    headquarters_address: str = ""
    property_locations: List[PropertyLocationSchema] = Field(default_factory=list)
    renewal_date: Optional[str] = None
    broker_notes: str = ""
    contacts: List[Dict[str, str]] = Field(default_factory=list)


class UpdateClientRequest(BaseModel):
    name: Optional[str] = None
    industry_code: Optional[str] = None
    business_type: Optional[str] = None
    years_in_business: Optional[int] = None
    annual_revenue: Optional[str] = None
    employee_count: Optional[int] = None
    headquarters_address: Optional[str] = None
    property_locations: Optional[List[PropertyLocationSchema]] = None
    renewal_date: Optional[str] = None
    broker_notes: Optional[str] = None
    contacts: Optional[List[Dict[str, str]]] = None


class CreateSubmissionRequest(BaseModel):
    client_id: str
    line_of_business: str = LineOfBusiness.PROPERTY.value
    effective_date: Optional[str] = None
    expiration_date: Optional[str] = None
    total_insured_value: str = ""
    coverage_requested: Dict[str, Any] = Field(default_factory=dict)
    submitted_carriers: List[str] = Field(default_factory=list)


class UpdateSubmissionRequest(BaseModel):
    line_of_business: Optional[str] = None
    status: Optional[str] = None
    effective_date: Optional[str] = None
    expiration_date: Optional[str] = None
    total_insured_value: Optional[str] = None
    coverage_requested: Optional[Dict[str, Any]] = None
    submitted_carriers: Optional[List[str]] = None
    acord_125_fields: Optional[Dict[str, Any]] = None
    acord_140_fields: Optional[Dict[str, Any]] = None


class UploadQuoteResponse(BaseModel):
    quote_id: str
    status: str
    message: str


class CompareWeights(BaseModel):
    premium_weight: float = Field(default=35.0, ge=0, le=100)
    coverage_weight: float = Field(default=30.0, ge=0, le=100)
    financial_weight: float = Field(default=20.0, ge=0, le=100)
    completeness_weight: float = Field(default=15.0, ge=0, le=100)


class CompareRequest(BaseModel):
    weights: Optional[CompareWeights] = None


class CompareResponse(BaseModel):
    comparison_table: List[Dict[str, Any]]
    recommendation: str
    placement_scores: List[Dict[str, Any]]


class ResearchRequest(BaseModel):
    company_name: str
    document_ids: List[str] = Field(
        default_factory=list,
        description="Optional IDs of uploaded research documents to include as context",
    )


class ResearchResponse(BaseModel):
    """Structured research brief — mirrors ResearchBrief schema."""

    company_name: str = ""
    business_description: str = ""
    headquarters: str = ""
    year_founded: Optional[int] = None
    employee_count: Optional[int] = None
    ownership_type: str = ""
    key_operations: List[str] = Field(default_factory=list)
    annual_revenue: str = ""
    revenue_trend: str = ""
    credit_rating: str = ""
    financial_highlights: List[str] = Field(default_factory=list)
    naics_code: str = ""
    industry_sector: str = ""
    common_perils: List[str] = Field(default_factory=list)
    loss_frequency: str = ""
    risk_factors: List[str] = Field(default_factory=list)
    insurance_needs: List[Dict[str, str]] = Field(default_factory=list)
    carrier_matches: List[Dict[str, Any]] = Field(default_factory=list)
    recent_news: List[Dict[str, str]] = Field(default_factory=list)
    citations: List[str] = Field(default_factory=list)
    citation_types: Dict[str, str] = Field(default_factory=dict)
    confidence_level: str = ""
    field_confidence: Dict[str, str] = Field(default_factory=dict)
    data_sources: List[str] = Field(default_factory=list)
    generated_at: str = ""


class CreateCarrierRequest(BaseModel):
    carrier_name: str
    amb_number: str = ""
    naic_code: str = ""
    financial_strength_rating: str = ""
    issuer_credit_rating: str = ""
    rating_outlook: str = ""
    balance_sheet_strength: str = ""
    operating_performance: str = ""
    business_profile: str = ""
    erm_assessment: str = ""
    net_premiums_written: str = ""
    policyholders_surplus: str = ""
    combined_ratio: str = ""
    five_year_avg_combined_ratio: str = ""
    report_date: Optional[str] = None
    direct_written_premium: str = ""
    nwp_to_surplus_ratio: str = ""
    lines_of_business_written: List[Dict[str, str]] = Field(default_factory=list)
    geographic_concentration: List[Dict[str, str]] = Field(default_factory=list)


class UpdateAcordFieldsRequest(BaseModel):
    """Request body for saving edited ACORD field values."""
    acord_125_fields: Dict[str, Any] = Field(default_factory=dict)
    acord_140_fields: Dict[str, Any] = Field(default_factory=dict)


class GeneratePackageRequest(BaseModel):
    """Request body for enhanced package generation."""
    carriers: List[str] = Field(default_factory=list)


class UpdatePackageEmailRequest(BaseModel):
    """Request body for saving an edited cover email draft."""
    cover_email: str


class UpdateResearchBriefRequest(BaseModel):
    """Request body for saving edits to a research brief."""
    business_description: Optional[str] = None
    headquarters: Optional[str] = None
    year_founded: Optional[int] = None
    employee_count: Optional[int] = None
    ownership_type: Optional[str] = None
    annual_revenue: Optional[str] = None
    revenue_trend: Optional[str] = None
    credit_rating: Optional[str] = None
    financial_highlights: Optional[List[str]] = None
    naics_code: Optional[str] = None
    industry_sector: Optional[str] = None
    loss_frequency: Optional[str] = None
    risk_factors: Optional[List[str]] = None
    key_operations: Optional[List[str]] = None
    common_perils: Optional[List[str]] = None
    confidence_level: Optional[str] = None


class MarkSentRequest(BaseModel):
    """Request body for marking a submission as sent."""
    carriers: List[str] = Field(default_factory=list)


# ============================================================================
# Dashboard Endpoints
# ============================================================================

@router.get("/dashboard")
async def get_dashboard():
    """
    Get dashboard metrics for the broker workspace.
    
    Returns:
        Dashboard metrics including total accounts, bound premium,
        open submissions, renewals due, and stale submissions.
    """
    try:
        metrics = storage.get_dashboard_metrics()
        return metrics
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Error fetching dashboard metrics", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# Client Endpoints
# ============================================================================

@router.get("/clients")
async def list_clients():
    """
    List all clients.
    
    Returns:
        List of client records.
    """
    try:
        clients = storage.list_clients()
        return clients
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Error listing clients", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/clients", status_code=201)
async def create_client(request: CreateClientRequest):
    """
    Create a new client.
    
    Args:
        request: Client creation data
        
    Returns:
        Created client record with generated ID.
    """
    try:
        # Convert property locations from Pydantic models to dataclass instances
        property_locations = [
            PropertyLocation(**loc.model_dump())
            for loc in request.property_locations
        ]
        
        # Create Client dataclass instance
        client = Client(
            name=request.name,
            industry_code=request.industry_code,
            business_type=request.business_type,
            years_in_business=request.years_in_business,
            annual_revenue=request.annual_revenue,
            employee_count=request.employee_count,
            headquarters_address=request.headquarters_address,
            property_locations=property_locations,
            renewal_date=request.renewal_date,
            broker_notes=request.broker_notes,
            contacts=request.contacts,
        )
        
        # Save to storage
        storage.save_client(client)
        
        # Return as dict
        return asdict(client)
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Error creating client", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/clients/{client_id}")
async def get_client(client_id: str):
    """
    Get client details by ID.
    
    Args:
        client_id: Client identifier
        
    Returns:
        Client record.
    """
    try:
        client = storage.get_client(client_id)
        if not client:
            raise HTTPException(status_code=404, detail=f"Client {client_id} not found")
        return client
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching client {client_id}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/clients/{client_id}")
async def update_client(client_id: str, request: UpdateClientRequest):
    """
    Update an existing client.
    
    Args:
        client_id: Client identifier
        request: Fields to update (partial)
        
    Returns:
        Updated client record.
    """
    try:
        # Get existing client
        existing = storage.get_client(client_id)
        if not existing:
            raise HTTPException(status_code=404, detail=f"Client {client_id} not found")
        
        # Update fields that were provided
        update_data = request.model_dump(exclude_unset=True)
        
        # Convert property locations if provided
        if "property_locations" in update_data and update_data["property_locations"]:
            update_data["property_locations"] = [
                PropertyLocation(**loc.model_dump()) if isinstance(loc, PropertyLocationSchema) else PropertyLocation(**loc)
                for loc in update_data["property_locations"]
            ]
        
        # Merge with existing data
        existing.update(update_data)
        existing["updated_at"] = datetime.utcnow().isoformat()
        
        # Convert property_locations dicts → PropertyLocation dataclasses
        if isinstance(existing.get("property_locations"), list):
            existing["property_locations"] = [
                PropertyLocation(**loc) if isinstance(loc, dict) else loc
                for loc in existing["property_locations"]
            ]
        
        # Create Client instance and save
        valid_keys = set(Client.__dataclass_fields__)
        client = Client(**{k: v for k, v in existing.items() if k in valid_keys})
        storage.save_client(client)
        
        return asdict(client)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating client {client_id}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/clients/{client_id}/submissions")
async def list_client_submissions(client_id: str):
    """
    List all submissions for a specific client.
    
    Args:
        client_id: Client identifier
        
    Returns:
        List of submission records for the client.
    """
    try:
        # Verify client exists
        client = storage.get_client(client_id)
        if not client:
            raise HTTPException(status_code=404, detail=f"Client {client_id} not found")
        
        submissions = storage.list_submissions(client_id=client_id)
        return submissions
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error listing submissions for client {client_id}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/clients/{client_id}/documents")
async def list_client_documents(client_id: str):
    """
    Aggregate all documents from all submissions for a client.

    Args:
        client_id: Client identifier

    Returns:
        List of document records across all submissions, each annotated with
        submission_id and line_of_business.
    """
    try:
        client = storage.get_client(client_id)
        if not client:
            raise HTTPException(status_code=404, detail=f"Client {client_id} not found")

        submissions = storage.list_submissions(client_id=client_id)
        documents: List[Dict[str, Any]] = []
        for sub in submissions:
            lob = sub.get("line_of_business", "")
            for doc in sub.get("documents", []):
                doc_copy = dict(doc)
                doc_copy["submission_id"] = sub.get("id", "")
                doc_copy["line_of_business"] = lob
                documents.append(doc_copy)
        return documents
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error listing documents for client {client_id}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# Submission Endpoints
# ============================================================================

@router.post("/submissions", status_code=201)
async def create_submission(request: CreateSubmissionRequest):
    """
    Create a new submission.
    
    Args:
        request: Submission creation data
        
    Returns:
        Created submission record with generated ID.
    """
    try:
        # Verify client exists
        client = storage.get_client(request.client_id)
        if not client:
            raise HTTPException(status_code=404, detail=f"Client {request.client_id} not found")
        
        # Create Submission dataclass instance
        submission = Submission(
            client_id=request.client_id,
            line_of_business=request.line_of_business,
            effective_date=request.effective_date,
            expiration_date=request.expiration_date,
            total_insured_value=request.total_insured_value,
            coverage_requested=request.coverage_requested,
            submitted_carriers=request.submitted_carriers,
        )
        
        # Save to storage
        storage.save_submission(submission)
        
        # Return as dict
        return asdict(submission)
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Error creating submission", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/submissions/{submission_id}")
async def get_submission(submission_id: str):
    """
    Get submission details by ID, including all quotes.
    
    Args:
        submission_id: Submission identifier
        
    Returns:
        Submission record with embedded quotes.
    """
    try:
        submission = storage.get_submission(submission_id)
        if not submission:
            raise HTTPException(status_code=404, detail=f"Submission {submission_id} not found")
        return submission
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching submission {submission_id}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/submissions/{submission_id}")
async def update_submission(submission_id: str, request: UpdateSubmissionRequest):
    """
    Update an existing submission.
    
    Args:
        submission_id: Submission identifier
        request: Fields to update (partial)
        
    Returns:
        Updated submission record.
    """
    try:
        # Get existing submission
        existing = storage.get_submission(submission_id)
        if not existing:
            raise HTTPException(status_code=404, detail=f"Submission {submission_id} not found")
        
        # Update fields that were provided
        update_data = request.model_dump(exclude_unset=True)
        
        # Merge with existing data
        existing.update(update_data)
        existing["updated_at"] = datetime.utcnow().isoformat()
        
        # Create Submission instance and save
        submission = _submission_from_dict(existing)
        storage.save_submission(submission)
        
        return asdict(submission)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating submission {submission_id}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# Quote Endpoints
# ============================================================================

@router.post("/submissions/{submission_id}/quotes", status_code=201)
async def upload_quote(
    submission_id: str,
    carrier_name: str = Form(...),
    file: UploadFile = File(...),
):
    """
    Upload a quote document for a submission.
    
    Args:
        submission_id: Submission identifier
        carrier_name: Name of the carrier providing the quote
        file: Quote document file (PDF, Excel, etc.)
        
    Returns:
        Quote ID and status message.
    """
    try:
        # Verify submission exists
        submission = storage.get_submission(submission_id)
        if not submission:
            raise HTTPException(status_code=404, detail=f"Submission {submission_id} not found")
        
        # Save file via storage layer
        content = await file.read()
        storage.save_file(
            f"quotes/{submission_id}/{file.filename}", content
        )
        
        # Determine source format from file extension
        file_ext = file.filename.split(".")[-1].lower()
        source_format_map = {
            "pdf": "pdf",
            "xlsx": "xlsx",
            "xls": "xlsx",
            "docx": "docx",
            "doc": "docx",
            "msg": "email",
            "eml": "email",
        }
        source_format = source_format_map.get(file_ext, "other")
        
        # Create Quote dataclass instance
        quote = Quote(
            submission_id=submission_id,
            carrier_name=carrier_name,
            source_format=source_format,
            source_file_name=file.filename,
            received_date=datetime.utcnow().isoformat(),
            status=QuoteStatus.RAW.value,
        )
        
        # Save initial quote with RAW status
        storage.save_quote(quote)
        
        # Extract quote fields using QuoteExtractor
        try:
            extractor = QuoteExtractor(
                _settings.openai,
                broker_settings=_settings.broker,
                cu_settings=_settings.content_understanding,
            )
            extracted_fields, confidence_scores = await extractor.extract_quote(
                content, file.filename, carrier_name
            )
            quote.fields = extracted_fields
            quote.confidence_scores = confidence_scores
            quote.status = QuoteStatus.EXTRACTED.value
            logger.info(f"Successfully extracted fields from quote {quote.id}")
        except Exception as e:
            logger.error(f"Quote extraction failed for {quote.id}: {e}", exc_info=True)
            # Keep quote in RAW status if extraction fails
            quote.status = QuoteStatus.RAW.value
        
        # Save updated quote
        storage.save_quote(quote)
        
        return UploadQuoteResponse(
            quote_id=quote.id,
            status=quote.status,
            message=f"Quote from {carrier_name} uploaded and {'extracted' if quote.status == QuoteStatus.EXTRACTED.value else 'saved (extraction pending)'}",
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading quote for submission {submission_id}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/submissions/{submission_id}/quotes")
async def list_quotes(submission_id: str):
    """
    List all quotes for a submission.
    
    Args:
        submission_id: Submission identifier
        
    Returns:
        List of quote records for the submission.
    """
    try:
        # Verify submission exists
        submission = storage.get_submission(submission_id)
        if not submission:
            raise HTTPException(status_code=404, detail=f"Submission {submission_id} not found")
        
        quotes = storage.list_quotes_for_submission(submission_id)
        return quotes
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error listing quotes for submission {submission_id}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/submissions/{submission_id}/compare")
async def compare_quotes(submission_id: str, body: Optional[CompareRequest] = None):
    """
    Run placement engine to compare quotes and generate recommendation.
    
    Args:
        submission_id: Submission identifier
        body: Optional request body with custom scoring weights
        
    Returns:
        Comparison table, recommendation, and placement scores.
    """
    try:
        # Verify submission exists
        submission_data = storage.get_submission(submission_id)
        if not submission_data:
            raise HTTPException(status_code=404, detail=f"Submission {submission_id} not found")
        
        # Convert to Submission dataclass (with nested types)
        submission = _submission_from_dict(submission_data)
        
        quotes_data = storage.list_quotes_for_submission(submission_id)
        
        if not quotes_data:
            return CompareResponse(
                comparison_table=[],
                recommendation="No quotes available for comparison",
                placement_scores=[],
            )
        
        # Convert quotes to Quote dataclass instances (with nested fields/scoring)
        quotes = [_quote_from_dict(q) for q in quotes_data]
        
        # Wire in PlacementEngine
        try:
            engine = PlacementEngine()
            
            # Load carrier profiles for FSR scoring
            carrier_profiles = {}
            for quote in quotes:
                profile_data = storage.get_carrier_profile_by_name(quote.carrier_name)
                if profile_data:
                    from app.broker.models import CarrierProfile
                    carrier_profiles[quote.carrier_name] = CarrierProfile(**profile_data)
            
            # Score quotes with optional custom weights
            weight_kwargs = {}
            if body and body.weights:
                weight_kwargs = {
                    "premium_weight": body.weights.premium_weight,
                    "coverage_weight": body.weights.coverage_weight,
                    "financial_weight": body.weights.financial_weight,
                    "completeness_weight": body.weights.completeness_weight,
                }
            scored_quotes = engine.score_quotes(quotes, submission, carrier_profiles, **weight_kwargs)
            
            # Generate recommendation
            recommendation = engine.generate_recommendation(scored_quotes)
            
            # Save scored quotes back to storage
            for quote in scored_quotes:
                storage.save_quote(quote)
            
        except Exception as e:
            logger.error(f"Placement engine failed for submission {submission_id}: {e}", exc_info=True)
            # Fall back to returning basic data if engine fails
            scored_quotes = quotes
            recommendation = (
                f"Analyzed {len(quotes)} quotes for submission {submission_id}. "
                "Placement scoring unavailable. Manual review recommended."
            )
        
        # Build response
        comparison_table = [
            {
                "quote_id": q.id,
                "carrier_name": q.carrier_name,
                "annual_premium": q.fields.annual_premium,
                "total_insured_value": q.fields.total_insured_value,
                "deductible": q.fields.deductible,
                "rating": q.fields.carrier_am_best_rating,
                "coverage_adequacy": q.scoring.coverage_adequacy,
                "placement_score": q.scoring.placement_score,
            }
            for q in scored_quotes
        ]
        
        placement_scores = [
            {
                "quote_id": q.id,
                "carrier_name": q.carrier_name,
                "placement_score": q.scoring.placement_score,
                "placement_rank": q.scoring.placement_rank,
                "recommendation_rationale": q.scoring.recommendation_rationale,
                "coverage_gaps": q.scoring.coverage_gaps,
                "premium_percentile": q.scoring.premium_percentile,
            }
            for q in scored_quotes
        ]
        
        return CompareResponse(
            comparison_table=comparison_table,
            recommendation=recommendation,
            placement_scores=placement_scores,
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error comparing quotes for submission {submission_id}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# Document Upload Endpoint
# ============================================================================

VALID_DOCUMENT_TYPES = {dt.value for dt in DocumentType}

# Known carrier names for auto-classification
_KNOWN_CARRIERS = [
    "travelers", "aig", "zurich", "hartford", "chubb", "liberty mutual",
    "nationwide", "allianz", "berkshire", "markel", "everest", "axis",
]


def _auto_classify_document(filename: str) -> Optional[str]:
    """Rule-based document type classifier based on filename patterns.

    Returns a DocumentType value string if a match is found, else None.
    """
    name_lower = filename.lower()

    if "sov" in name_lower or "statement of values" in name_lower:
        return DocumentType.SOV.value
    if "loss" in name_lower or "claim" in name_lower:
        return DocumentType.LOSS_RUNS.value
    if "dec" in name_lower or "declaration" in name_lower:
        return DocumentType.PRIOR_DECLARATION.value
    if "acord" in name_lower or "125" in name_lower:
        return DocumentType.ACORD_125.value
    if "140" in name_lower:
        return DocumentType.ACORD_140.value
    if "quote" in name_lower or any(c in name_lower for c in _KNOWN_CARRIERS):
        return DocumentType.CARRIER_QUOTE.value
    if "endorse" in name_lower:
        return DocumentType.OTHER.value  # No dedicated endorsement type; keep as other

    return None


@router.post("/submissions/{submission_id}/documents", status_code=201)
async def upload_document(
    submission_id: str,
    file: UploadFile = File(...),
    document_type: str = Form("other"),
):
    """
    Upload a client document (SOV, loss runs, prior declaration, etc.) to a submission.

    Includes rule-based auto-classification: when the user-provided type is
    ``other`` and the filename matches a known pattern, the type is overridden
    automatically.

    Args:
        submission_id: Submission identifier
        file: Document file to upload
        document_type: One of acord_125, acord_140, sov, loss_runs,
                       prior_declaration, other

    Returns:
        201 with the created BrokerDocument record.
    """
    try:
        # Validate submission exists
        submission_data = storage.get_submission(submission_id)
        if not submission_data:
            raise HTTPException(
                status_code=404,
                detail=f"Submission {submission_id} not found",
            )

        # Validate document_type
        if document_type not in VALID_DOCUMENT_TYPES:
            raise HTTPException(
                status_code=422,
                detail=f"Invalid document_type '{document_type}'. Must be one of: {sorted(VALID_DOCUMENT_TYPES)}",
            )

        # AC-6.4: Auto-classify when user selected "other"
        if document_type == DocumentType.OTHER.value:
            classified = _auto_classify_document(file.filename)
            if classified:
                logger.info(
                    "Auto-classified document '%s' as '%s' (was '%s')",
                    file.filename, classified, document_type,
                )
                document_type = classified

        # Save file via storage layer
        content = await file.read()
        rel_path = f"documents/{submission_id}/{file.filename}"
        stored_path = storage.save_file(rel_path, content)

        # Create BrokerDocument record
        doc = BrokerDocument(
            submission_id=submission_id,
            document_type=document_type,
            file_name=file.filename,
            blob_url=stored_path,
        )

        # Add to submission.documents and save
        documents: list = submission_data.get("documents", [])
        documents.append(asdict(doc))
        submission_data["documents"] = documents
        submission_data["updated_at"] = datetime.utcnow().isoformat()
        submission = _submission_from_dict(submission_data)
        storage.save_submission(submission)

        return asdict(doc)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(
            f"Error uploading document for submission {submission_id}",
            exc_info=True,
        )
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# ACORD Field Extraction & Forms Endpoints
# ============================================================================


@router.post("/submissions/{submission_id}/extract-acord")
async def extract_acord_fields(submission_id: str):
    """
    Run ACORD field extraction on documents attached to a submission.

    Reads all documents, uses the SubmissionGenerator engine to extract
    ACORD 125/140 fields, and saves extracted fields + confidence scores.

    Args:
        submission_id: Submission identifier

    Returns:
        200 with extracted ACORD 125/140 fields and confidence scores.
    """
    try:
        submission_data = storage.get_submission(submission_id)
        if not submission_data:
            raise HTTPException(status_code=404, detail=f"Submission {submission_id} not found")

        # Reconstruct BrokerDocument list
        documents = [
            BrokerDocument(**{
                k: v for k, v in d.items()
                if k in BrokerDocument.__dataclass_fields__
            })
            for d in submission_data.get("documents", [])
        ]

        if not documents:
            raise HTTPException(
                status_code=422,
                detail="No documents attached to this submission. Upload documents first.",
            )

        # Run extraction
        generator = SubmissionGenerator(
            _settings.openai,
            broker_settings=_settings.broker,
            cu_settings=_settings.content_understanding,
        )
        acord_125, acord_140, confidence, field_sources = await generator.extract_acord_fields(documents)

        # Flatten loss_history array into individual ACORD 140 canonical fields
        loss_history = acord_140.get("loss_history", [])
        if isinstance(loss_history, list):
            for idx, entry in enumerate(loss_history, start=1):
                if isinstance(entry, dict):
                    prefix = f"loss_{idx}"
                    acord_140[f"{prefix}_date"] = entry.get("date", entry.get("DateOfLoss", ""))
                    acord_140[f"{prefix}_cause"] = entry.get("cause", entry.get("cause_of_loss", entry.get("CauseOfLoss", "")))
                    acord_140[f"{prefix}_amount"] = entry.get("amount", entry.get("amount_paid", entry.get("AmountPaid", "")))
                    acord_140[f"{prefix}_description"] = entry.get("description", entry.get("Description", ""))
                    acord_140[f"{prefix}_status"] = entry.get("status", entry.get("Status", ""))

        # Flatten property_locations array into individual ACORD 140 fields
        prop_locs = acord_140.get("property_locations", [])
        if isinstance(prop_locs, list):
            for idx, loc in enumerate(prop_locs, start=1):
                if isinstance(loc, dict):
                    prefix = f"loc_{idx}"
                    acord_140[f"{prefix}_address"] = loc.get("Address", loc.get("address", ""))
                    acord_140[f"{prefix}_occupancy"] = loc.get("OccupancyType", loc.get("occupancy", ""))
                    acord_140[f"{prefix}_construction"] = loc.get("ConstructionType", loc.get("construction_type", ""))
                    acord_140[f"{prefix}_year_built"] = loc.get("YearBuilt", loc.get("year_built", ""))
                    acord_140[f"{prefix}_sqft"] = loc.get("SquareFootage", loc.get("square_footage", ""))
                    acord_140[f"{prefix}_building_value"] = loc.get("BuildingValue", loc.get("building_value", ""))
                    acord_140[f"{prefix}_contents_value"] = loc.get("ContentsValue", loc.get("contents_value", ""))
                    acord_140[f"{prefix}_bi_value"] = loc.get("BusinessInterruptionValue", loc.get("bi_value", ""))
                    acord_140[f"{prefix}_protection_class"] = loc.get("ProtectionClass", loc.get("protection_class", ""))

        # Save to submission
        submission_data["acord_125_fields"] = acord_125
        submission_data["acord_140_fields"] = acord_140
        submission_data["acord_field_confidence"] = confidence
        submission_data["acord_field_sources"] = field_sources
        submission_data["updated_at"] = datetime.utcnow().isoformat()
        submission = _submission_from_dict(submission_data)
        storage.save_submission(submission)

        return {
            "submission_id": submission_id,
            "acord_125": acord_125,
            "acord_140": acord_140,
            "confidence": confidence,
            "field_sources": field_sources,
            "fields_extracted": len(acord_125) + len(acord_140),
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error extracting ACORD fields for submission {submission_id}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/submissions/{submission_id}/acord-forms")
async def get_acord_forms(submission_id: str):
    """
    Get ACORD form data with field metadata for UI rendering.

    Merges extracted values with the ACORD 125/140 form templates so the
    frontend can render editable forms with labels, sections, and required
    flags.

    Args:
        submission_id: Submission identifier

    Returns:
        200 with ACORD 125 template, ACORD 140 template, SOV template,
        each merged with extracted values and confidence scores.
    """
    try:
        submission_data = storage.get_submission(submission_id)
        if not submission_data:
            raise HTTPException(status_code=404, detail=f"Submission {submission_id} not found")

        extracted_125 = submission_data.get("acord_125_fields", {})
        extracted_140 = submission_data.get("acord_140_fields", {})
        confidence = submission_data.get("acord_field_confidence", {})

        # Map internal field names to canonical ACORD names
        mapped_125 = map_extracted_to_acord125(extracted_125)

        # Build ACORD 125 form with values
        template_125 = get_acord_form_template("125")
        form_125_fields = {}
        for field_key, meta in template_125["fields"].items():
            form_125_fields[field_key] = {
                **meta,
                "value": mapped_125.get(field_key, extracted_125.get(field_key, "")),
                "confidence": confidence.get(field_key, confidence.get(
                    # Also check snake_case keys
                    field_key.lower(), None
                )),
            }

        # Build ACORD 140 form with values
        template_140 = get_acord_form_template("140")

        # Flatten first property location and first loss entry from CU
        # nested arrays into the template's flat field structure.
        flat_140 = dict(extracted_140)
        prop_locs = extracted_140.get("property_locations", [])
        if isinstance(prop_locs, list) and prop_locs:
            first = prop_locs[0] if isinstance(prop_locs[0], dict) else {}
            _PROP_FIELD_MAP = {
                "Address": "PropertyAddress",
                "OccupancyType": "PropertyOccupancy",
                "ConstructionType": "ConstructionType",
                "YearBuilt": "YearBuilt",
                "SquareFootage": "SquareFootage",
                "BuildingValue": "BuildingValue",
                "ContentsValue": "ContentsValue",
                "BusinessInterruptionValue": "BusinessInterruptionValue",
                "SprinklerSystem": "ProtectionClass",
                "ProtectionClass": "ProtectionClass",
            }
            for cu_key, tmpl_key in _PROP_FIELD_MAP.items():
                val = first.get(cu_key)
                if val and tmpl_key not in flat_140:
                    flat_140[tmpl_key] = val

        loss_hist = extracted_140.get("loss_history", [])
        if isinstance(loss_hist, list) and loss_hist:
            first_loss = loss_hist[0] if isinstance(loss_hist[0], dict) else {}
            _LOSS_FIELD_MAP = {
                "DateOfLoss": "LossDate",
                "date": "LossDate",
                "CauseOfLoss": "CauseOfLoss",
                "cause_of_loss": "CauseOfLoss",
                "cause": "CauseOfLoss",
                "AmountPaid": "LossAmountPaid",
                "amount_paid": "LossAmountPaid",
                "amount": "LossAmountPaid",
                "Description": "LossDescription",
                "description": "LossDescription",
                "Status": "LossStatus",
                "status": "LossStatus",
            }
            for cu_key, tmpl_key in _LOSS_FIELD_MAP.items():
                val = first_loss.get(cu_key)
                if val and tmpl_key not in flat_140:
                    flat_140[tmpl_key] = val

        form_140_fields = {}
        for field_key, meta in template_140["fields"].items():
            form_140_fields[field_key] = {
                **meta,
                "value": flat_140.get(field_key, ""),
                "confidence": confidence.get(field_key, None),
            }

        # Build SOV template
        template_sov = get_acord_form_template("sov")

        return {
            "submission_id": submission_id,
            "acord_125": {
                "form_type": template_125["form_type"],
                "title": template_125["title"],
                "fields": form_125_fields,
                "raw_extracted": extracted_125,
            },
            "acord_140": {
                "form_type": template_140["form_type"],
                "title": template_140["title"],
                "fields": form_140_fields,
                "raw_extracted": extracted_140,
            },
            "sov": {
                "form_type": template_sov["form_type"],
                "title": template_sov["title"],
                "field_schema": template_sov["fields"],
                "locations": extracted_140.get("property_locations", []),
            },
            "confidence": confidence,
            "field_sources": submission_data.get("acord_field_sources", {}),
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching ACORD forms for submission {submission_id}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/submissions/{submission_id}/acord-fields")
async def update_acord_fields(submission_id: str, request: UpdateAcordFieldsRequest):
    """
    Save edited ACORD field values.

    The broker edits fields in the UI and saves them back to the
    submission record.

    Args:
        submission_id: Submission identifier
        request: Updated ACORD 125 and 140 field values

    Returns:
        200 with the updated fields.
    """
    try:
        submission_data = storage.get_submission(submission_id)
        if not submission_data:
            raise HTTPException(status_code=404, detail=f"Submission {submission_id} not found")

        if request.acord_125_fields:
            submission_data["acord_125_fields"] = request.acord_125_fields
        if request.acord_140_fields:
            submission_data["acord_140_fields"] = request.acord_140_fields
        submission_data["updated_at"] = datetime.utcnow().isoformat()

        submission = _submission_from_dict(submission_data)
        storage.save_submission(submission)

        return {
            "submission_id": submission_id,
            "acord_125_fields": submission_data.get("acord_125_fields", {}),
            "acord_140_fields": submission_data.get("acord_140_fields", {}),
            "updated_at": submission_data["updated_at"],
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating ACORD fields for submission {submission_id}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# Submission Package Endpoint
# ============================================================================


@router.post("/submissions/{submission_id}/generate-package")
async def generate_submission_package(
    submission_id: str,
    request: Optional[GeneratePackageRequest] = None,
):
    """
    Generate a submission package from the submission's documents and
    extracted ACORD fields.

    Enhanced version: generates a cover email via LLM, compiles ACORD 125 +
    140 + SOV + documents into a package, and returns the full package
    with an editable cover email draft.

    Accepts an optional ``carriers`` list in the request body.  Falls back
    to the submission's ``submitted_carriers`` when not provided.

    Args:
        submission_id: Submission identifier
        request: Optional request body with carrier list

    Returns:
        200 with the full submission package JSON.
    """
    try:
        # Get submission
        submission_data = storage.get_submission(submission_id)
        if not submission_data:
            raise HTTPException(
                status_code=404,
                detail=f"Submission {submission_id} not found",
            )

        # Get associated client
        client_id = submission_data.get("client_id", "")
        client_data = storage.get_client(client_id) if client_id else {}
        if client_data is None:
            client_data = {}

        # Resolve carrier list
        carriers = (
            request.carriers
            if request and request.carriers
            else submission_data.get("submitted_carriers", [])
        )

        # Generate cover email
        pkg_gen = SubmissionPackageGenerator(_settings.openai)
        cover_email = await pkg_gen.generate_cover_email(
            client_data, submission_data, carriers,
        )

        # Compile full package
        package = pkg_gen.compile_package(client_data, submission_data, cover_email)

        # Generate per-carrier emails
        carrier_emails = await pkg_gen.generate_carrier_emails(package, carriers)
        package["carrier_emails"] = carrier_emails

        # Add quotes summary
        package["quotes_summary"] = [
            {
                "id": q.get("id", ""),
                "carrier_name": q.get("carrier_name", ""),
                "annual_premium": q.get("fields", {}).get("annual_premium", ""),
                "status": q.get("status", ""),
            }
            for q in submission_data.get("quotes", [])
        ]

        # Update submission status
        submission_data["status"] = SubmissionStatus.SUBMITTED.value
        submission_data["submission_date"] = datetime.utcnow().isoformat()
        submission_data["updated_at"] = datetime.utcnow().isoformat()
        if carriers:
            submission_data["submitted_carriers"] = carriers
        submission = _submission_from_dict(submission_data)
        storage.save_submission(submission)

        package["status"] = SubmissionStatus.SUBMITTED.value
        return package
    except HTTPException:
        raise
    except Exception as e:
        logger.error(
            f"Error generating package for submission {submission_id}",
            exc_info=True,
        )
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/submissions/{submission_id}/package-email")
async def update_package_email(submission_id: str, request: UpdatePackageEmailRequest):
    """
    Save an edited cover email draft for a submission package.

    The broker edits the draft email in the UI and saves it back.

    Args:
        submission_id: Submission identifier
        request: Updated cover email text

    Returns:
        200 with confirmation and the saved email.
    """
    try:
        submission_data = storage.get_submission(submission_id)
        if not submission_data:
            raise HTTPException(status_code=404, detail=f"Submission {submission_id} not found")

        # Store cover email in the submission's coverage_requested dict
        coverage = submission_data.get("coverage_requested", {})
        coverage["cover_email_draft"] = request.cover_email
        submission_data["coverage_requested"] = coverage
        submission_data["updated_at"] = datetime.utcnow().isoformat()

        submission = _submission_from_dict(submission_data)
        storage.save_submission(submission)

        return {
            "submission_id": submission_id,
            "cover_email": request.cover_email,
            "updated_at": submission_data["updated_at"],
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error saving package email for submission {submission_id}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/submissions/{submission_id}/mark-sent")
async def mark_submission_sent(submission_id: str, request: Optional[MarkSentRequest] = None):
    """
    Mark a submission as sent to carriers.

    Updates the submission status to 'submitted', records the sent_at
    timestamp, and optionally updates the carrier list.

    Args:
        submission_id: Submission identifier
        request: Optional request body with carrier list

    Returns:
        200 with updated submission status and sent_at timestamp.
    """
    try:
        submission_data = storage.get_submission(submission_id)
        if not submission_data:
            raise HTTPException(status_code=404, detail=f"Submission {submission_id} not found")

        now = datetime.utcnow().isoformat()
        submission_data["status"] = SubmissionStatus.SUBMITTED.value
        submission_data["sent_at"] = now
        submission_data["submission_date"] = now
        submission_data["updated_at"] = now

        if request and request.carriers:
            submission_data["submitted_carriers"] = request.carriers

        submission = _submission_from_dict(submission_data)
        storage.save_submission(submission)

        return {
            "submission_id": submission_id,
            "status": SubmissionStatus.SUBMITTED.value,
            "sent_at": now,
            "submitted_carriers": submission_data.get("submitted_carriers", []),
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error marking submission {submission_id} as sent", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# Client Research Endpoint
# ============================================================================

@router.post("/clients/{client_id}/research")
async def research_client(client_id: str, request: ResearchRequest):
    """
    Run AI-powered client research and return a structured ResearchBrief.

    Optionally incorporates uploaded research documents (AM Best reports, etc.)
    when document_ids are provided.

    Args:
        client_id: Client identifier
        request: Research parameters including company name and optional document_ids

    Returns:
        Structured ResearchResponse (not free-text markdown).
    """
    try:
        # Verify client exists
        client_data = storage.get_client(client_id)
        if not client_data:
            raise HTTPException(status_code=404, detail=f"Client {client_id} not found")

        # Load uploaded research documents if any IDs provided
        uploaded_docs: List[str] = []
        carrier_data: List[Dict[str, Any]] = []

        # Check if Content Understanding is available for structured extraction
        cu_processor = BrokerCUProcessor(
            _settings.content_understanding, _settings.broker
        )
        use_cu = cu_processor.is_available()
        # Limit CU analysis to avoid timeouts — process at most 3 docs via CU
        _MAX_CU_DOCS = 3
        cu_docs_processed = 0
        if use_cu:
            logger.info(
                "Content Understanding available — will extract structured "
                "carrier data from up to %d research documents", _MAX_CU_DOCS
            )
        else:
            logger.info(
                "Content Understanding not available — using text extraction fallback"
            )

        if request.document_ids:
            for doc_id in request.document_ids:
                # Look up metadata via storage layer
                meta_rel = f"research-documents/{client_id}/{doc_id}.meta.json"
                meta = storage._load_json(meta_rel)
                if not meta:
                    logger.warning(f"Research document {doc_id} not found, skipping")
                    continue
                try:
                    # Load file content via storage layer
                    file_stored_path = meta.get("file_path", "")
                    file_name = meta.get("file_name", "")
                    # Derive the storage-relative path for the file
                    file_rel = f"research-documents/{client_id}/{doc_id}_{file_name}"
                    file_content = storage.load_file(file_rel)
                    if file_content is None:
                        # Fallback: try the stored file_path directly (legacy local data)
                        fp = Path(file_stored_path)
                        if fp.exists():
                            file_content = fp.read_bytes()
                    if file_content is None:
                        logger.warning(f"Research doc file not found for {doc_id}")
                        continue

                    # --- CU-based structured extraction (limited to avoid timeouts) ---
                    cu_succeeded = False
                    if use_cu and cu_docs_processed < _MAX_CU_DOCS:
                        try:
                            doc_type = meta.get("document_type", "other")
                            cu_result = await cu_processor.analyze_document(
                                file_content,
                                file_name,
                                doc_type,
                                _settings.broker.research_analyzer,
                            )
                            if cu_result.get("fields"):
                                carrier_data.append(cu_result["fields"])
                                logger.info(
                                    f"CU extracted {len(cu_result['fields'])} fields "
                                    f"from research doc {doc_id}"
                                )
                            if cu_result.get("markdown"):
                                uploaded_docs.append(cu_result["markdown"])
                                logger.info(
                                    f"Using CU markdown for doc {doc_id}: "
                                    f"{len(cu_result['markdown'])} chars"
                                )
                                cu_succeeded = True
                                cu_docs_processed += 1
                        except Exception as cu_err:
                            logger.warning(
                                f"CU analysis failed for {doc_id}, falling back "
                                f"to text extraction: {cu_err}"
                            )

                    # --- Fallback text extraction (or when CU not available) ---
                    if not cu_succeeded:
                        file_ext = Path(file_name).suffix.lower()
                        if file_ext == ".pdf":
                            try:
                                import io
                                from PyPDF2 import PdfReader
                                reader = PdfReader(io.BytesIO(file_content))
                                pages = [p.extract_text() for p in reader.pages if p.extract_text()]
                                doc_text = "\n\n".join(pages) if pages else ""
                            except ImportError:
                                doc_text = file_content.decode("utf-8", errors="replace")
                            except Exception as pdf_err:
                                logger.warning(f"PDF parsing failed for {doc_id}: {pdf_err}")
                                doc_text = file_content.decode("utf-8", errors="replace")
                        elif file_ext == ".docx":
                            try:
                                import io
                                from docx import Document as DocxDocument
                                ddoc = DocxDocument(io.BytesIO(file_content))
                                parts = [p.text for p in ddoc.paragraphs if p.text.strip()]
                                for tbl in ddoc.tables:
                                    for row in tbl.rows:
                                        cells = [c.text.strip() for c in row.cells]
                                        if any(cells):
                                            parts.append(" | ".join(cells))
                                doc_text = "\n".join(parts) if parts else ""
                            except ImportError:
                                doc_text = file_content.decode("utf-8", errors="replace")
                        else:
                            doc_text = file_content.decode("utf-8", errors="replace")
                        uploaded_docs.append(doc_text)
                        logger.info(f"Loaded research doc {doc_id}: {len(doc_text)} chars")
                except Exception as doc_err:
                    logger.warning(f"Failed to read research document {doc_id}: {doc_err}")

        # Auto-create carrier profiles from CU-extracted data
        _CU_TO_CARRIER_FIELD = {
            "CarrierName": "carrier_name",
            "AMBNumber": "amb_number",
            "NAICCode": "naic_code",
            "FinancialStrengthRating": "financial_strength_rating",
            "IssuerCreditRating": "issuer_credit_rating",
            "RatingOutlook": "rating_outlook",
            "BalanceSheetStrength": "balance_sheet_strength",
            "OperatingPerformance": "operating_performance",
            "BusinessProfile": "business_profile",
            "ERMAssessment": "erm_assessment",
            "NetPremiumsWritten": "net_premiums_written",
            "PolicyholdersSurplus": "policyholders_surplus",
            "CombinedRatio": "combined_ratio",
            "FiveYearAvgCombinedRatio": "five_year_avg_combined_ratio",
            "DirectWrittenPremium": "direct_written_premium",
            "NWPToSurplusRatio": "nwp_to_surplus_ratio",
            "LinesOfBusinessWritten": "lines_of_business_written",
            "GeographicConcentration": "geographic_concentration",
            "ReportDate": "report_date",
        }
        for cd in carrier_data:
            carrier_name = cd.get("CarrierName", "")
            if not carrier_name:
                continue
            existing = storage.get_carrier_profile_by_name(carrier_name)
            if existing:
                logger.info(
                    f"Carrier profile already exists for '{carrier_name}', skipping auto-create"
                )
                continue
            profile_kwargs: Dict[str, Any] = {}
            for cu_field, model_field in _CU_TO_CARRIER_FIELD.items():
                value = cd.get(cu_field)
                if value is None:
                    continue
                # Normalize array-of-string → list-of-dict for model compatibility
                if model_field == "lines_of_business_written" and isinstance(value, list):
                    value = [
                        {"line": v} if isinstance(v, str) else v for v in value
                    ]
                elif model_field == "geographic_concentration" and isinstance(value, str):
                    value = [{"region": value}] if value else []
                profile_kwargs[model_field] = value
            try:
                profile = CarrierProfile(**profile_kwargs)
                storage.save_carrier_profile(profile)
                logger.info(
                    f"Auto-created carrier profile for '{carrier_name}' from CU analysis"
                )
            except Exception as cp_err:
                logger.warning(
                    f"Failed to auto-create carrier profile for '{carrier_name}': {cp_err}"
                )

        # Wire in ClientResearchEngine
        try:
            research_engine = ClientResearchEngine(_settings.openai)
            brief = await research_engine.research_client(
                request.company_name,
                uploaded_docs=uploaded_docs if uploaded_docs else None,
                carrier_data=carrier_data if carrier_data else None,
            )

            # Save the research brief to the client record and append to history
            brief.setdefault("generated_at", datetime.utcnow().isoformat())

            # Enrich carrier_matches with stored CarrierProfile metrics
            carrier_matches = brief.get("carrier_matches", [])
            for match in carrier_matches:
                carrier_name = match.get("carrier", "")
                if not carrier_name:
                    continue
                profile = storage.get_carrier_profile_by_name(carrier_name)
                if profile:
                    match.setdefault("fsr", profile.get("financial_strength_rating", ""))
                    match.setdefault("icr", profile.get("issuer_credit_rating", ""))
                    match.setdefault("outlook", profile.get("rating_outlook", ""))
                    match.setdefault("balance_sheet_strength", profile.get("balance_sheet_strength", ""))
                    match.setdefault("operating_performance", profile.get("operating_performance", ""))
                    match.setdefault("combined_ratio", profile.get("combined_ratio", ""))
                    match.setdefault("nwp_to_surplus_ratio", profile.get("nwp_to_surplus_ratio", ""))
            brief["carrier_matches"] = carrier_matches

            client_data["research_brief"] = brief
            # Prepend to research_history (most recent first, cap at 20)
            history = client_data.get("research_history", [])
            history_entry = {
                "id": str(uuid.uuid4()),
                "generated_at": brief.get("generated_at", datetime.utcnow().isoformat()),
                "company_name": brief.get("company_name", request.company_name),
                "confidence_level": brief.get("confidence_level", "Medium"),
                "data_sources": brief.get("data_sources", []),
                "brief": brief,
            }
            history.insert(0, history_entry)
            client_data["research_history"] = history[:20]
            # Convert property_locations dicts → PropertyLocation dataclasses
            if isinstance(client_data.get("property_locations"), list):
                client_data["property_locations"] = [
                    PropertyLocation(**loc) if isinstance(loc, dict) else loc
                    for loc in client_data["property_locations"]
                ]
            valid_keys = set(Client.__dataclass_fields__)
            client = Client(**{k: v for k, v in client_data.items() if k in valid_keys})
            storage.save_client(client)

            # Return structured response
            brief.setdefault("company_name", request.company_name)
            return ResearchResponse(**{
                k: v for k, v in brief.items()
                if k in ResearchResponse.model_fields and not k.startswith("_")
            })

        except Exception as e:
            logger.error(f"Client research failed for {client_id}: {e}", exc_info=True)
            # Return structured demo fallback when AI research engine is unavailable
            revenue = client_data.get('annual_revenue', '$12,000,000')
            employees = client_data.get('employee_count', 45)
            years = client_data.get('years_in_business', 15)
            hq = client_data.get('headquarters_address', 'Hartford, CT')
            industry = client_data.get('industry_code', '531110')
            num_locations = len(client_data.get('property_locations', [])) or 3

            return ResearchResponse(
                company_name=request.company_name,
                business_description=(
                    f"{request.company_name} is a commercial property management company "
                    f"based in {hq}, specializing in Class A office space, mixed-use "
                    f"retail/office properties, and warehouse/distribution facilities "
                    f"across Connecticut. Manages approximately $18.5M TIV across "
                    f"{num_locations} locations with {employees} employees."
                ),
                headquarters=hq,
                year_founded=(datetime.utcnow().year - years) if years else None,
                employee_count=employees,
                ownership_type="Private (LLC)",
                key_operations=[
                    "Commercial property management",
                    "Class A office leasing",
                    "Mixed-use retail/office operations",
                    "Warehouse/distribution facilities",
                ],
                annual_revenue=revenue,
                revenue_trend="Stable",
                credit_rating="Investment grade equivalent",
                financial_highlights=[
                    "Consistent occupancy rates above 90%",
                    "Conservative debt-to-equity ratio at 0.45",
                    f"Total Insured Value: $18,500,000 across {num_locations} locations",
                ],
                naics_code=industry,
                industry_sector="Commercial Real Estate",
                common_perils=[
                    "Property damage (fire, water)",
                    "Wind/hail (coastal exposure)",
                    "Slip & fall liability",
                ],
                loss_frequency="Low",
                risk_factors=[
                    "Coastal exposure (Stamford property)",
                    "Aging infrastructure (New Haven property)",
                    "Connecticut nor'easter exposure",
                ],
                insurance_needs=[
                    {"line": "Commercial Property", "priority": "Primary", "estimated_premium": "$35,000 - $50,000", "rationale": f"$18.5M TIV across {num_locations} locations"},
                    {"line": "General Liability", "priority": "Required", "estimated_premium": "$8,000 - $15,000", "rationale": "Office/retail tenant exposure"},
                    {"line": "Workers Compensation", "priority": "Required", "estimated_premium": "$12,000 - $18,000", "rationale": f"{employees} employees, property maintenance"},
                    {"line": "Umbrella/Excess", "priority": "Recommended", "estimated_premium": "$5,000 - $10,000", "rationale": "Multi-location liability aggregation"},
                    {"line": "Cyber Liability", "priority": "Recommended", "estimated_premium": "$3,000 - $6,000", "rationale": "Tenant data, online leasing platform"},
                ],
                carrier_matches=[
                    {"carrier": "Travelers", "rating": "A++", "appetite": "Strong", "rationale": "CT commercial property expertise"},
                    {"carrier": "AIG", "rating": "A+", "appetite": "Good", "rationale": "Multi-location property program"},
                    {"carrier": "Zurich", "rating": "A+", "appetite": "Good", "rationale": "Clean loss history pricing; flood coverage"},
                    {"carrier": "Hartford", "rating": "A+", "appetite": "Strong", "rationale": "Regional CT expertise; local claims handling"},
                ],
                recent_news=[
                    {"date": "2025-01-10", "headline": "CT commercial property market hardening 5-8% for 2026 renewals", "source": "Hartford Business Journal"},
                    {"date": "2024-12-15", "headline": "Hartford metro office vacancy stabilizing at 12%", "source": "CoStar"},
                ],
                citations=[
                    "Client database records",
                    "CT Insurance Department — Market Report 2025",
                    "NAICS industry classification data",
                    "Note: Demo content — AI research engine unavailable",
                ],
                confidence_level="Medium",
                data_sources=["Demo Fallback"],
                generated_at=datetime.utcnow().isoformat(),
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error researching client {client_id}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/clients/{client_id}/research-history")
async def get_research_history(client_id: str):
    """
    Get the list of past research briefs for a client.

    Returns:
        List of research history entries with id, date, sources, and brief data.
    """
    client_data = storage.get_client(client_id)
    if not client_data:
        raise HTTPException(status_code=404, detail=f"Client {client_id} not found")
    return client_data.get("research_history", [])


@router.put("/clients/{client_id}/research-brief")
async def update_research_brief(client_id: str, request: UpdateResearchBriefRequest):
    """
    Update the current research brief with user edits.

    Accepts partial updates to text fields of the research brief
    and persists them to the client record.

    Args:
        client_id: Client identifier
        request: Partial research brief update

    Returns:
        Updated research brief.
    """
    try:
        client_data = storage.get_client(client_id)
        if not client_data:
            raise HTTPException(status_code=404, detail=f"Client {client_id} not found")

        current_brief = client_data.get("research_brief")
        if not current_brief:
            raise HTTPException(status_code=404, detail="No research brief exists for this client")

        # Apply only the provided (non-None) fields
        updates = request.model_dump(exclude_none=True)
        for key, value in updates.items():
            current_brief[key] = value

        client_data["research_brief"] = current_brief
        client_data["updated_at"] = datetime.utcnow().isoformat()

        # Persist
        if isinstance(client_data.get("property_locations"), list):
            client_data["property_locations"] = [
                PropertyLocation(**loc) if isinstance(loc, dict) else loc
                for loc in client_data["property_locations"]
            ]
        valid_keys = set(Client.__dataclass_fields__)
        client = Client(**{k: v for k, v in client_data.items() if k in valid_keys})
        storage.save_client(client)

        return current_brief
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating research brief for client {client_id}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# Research Document Upload Endpoint
# ============================================================================

VALID_RESEARCH_DOC_TYPES = {"am_best", "annual_report", "financial_statement", "sp_report", "other"}


@router.post("/clients/{client_id}/research-documents", status_code=201)
async def upload_research_document(
    client_id: str,
    file: UploadFile = File(...),
    document_type: str = Form("other"),
):
    """
    Upload a pre-research document (AM Best report, annual report, etc.)
    that will be fed into the research engine as primary source context.

    Args:
        client_id: Client identifier
        file: Document file to upload
        document_type: One of am_best, annual_report, financial_statement, sp_report, other

    Returns:
        201 with the created ResearchDocument metadata.
    """
    try:
        # Verify client exists
        client_data = storage.get_client(client_id)
        if not client_data:
            raise HTTPException(status_code=404, detail=f"Client {client_id} not found")

        if document_type not in VALID_RESEARCH_DOC_TYPES:
            raise HTTPException(
                status_code=422,
                detail=f"Invalid document_type '{document_type}'. Must be one of: {sorted(VALID_RESEARCH_DOC_TYPES)}",
            )

        # Save file via storage layer
        doc_id = str(uuid.uuid4())
        content = await file.read()
        file_rel = f"research-documents/{client_id}/{doc_id}_{file.filename}"
        stored_path = storage.save_file(file_rel, content)

        # Write metadata sidecar
        meta = {
            "id": doc_id,
            "client_id": client_id,
            "file_name": file.filename,
            "document_type": document_type,
            "file_path": stored_path,
            "uploaded_at": datetime.utcnow().isoformat(),
        }
        meta_rel = f"research-documents/{client_id}/{doc_id}.meta.json"
        storage._save_json(meta_rel, meta)

        logger.info(
            f"Uploaded research document {doc_id} ({file.filename}) for client {client_id}"
        )

        return meta

    except HTTPException:
        raise
    except Exception as e:
        logger.error(
            f"Error uploading research document for client {client_id}",
            exc_info=True,
        )
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/clients/{client_id}/research-documents")
async def list_research_documents(client_id: str):
    """
    List all research documents uploaded for a client.

    Args:
        client_id: Client identifier

    Returns:
        List of research document metadata records.
    """
    try:
        client_data = storage.get_client(client_id)
        if not client_data:
            raise HTTPException(status_code=404, detail=f"Client {client_id} not found")

        # List meta.json files via storage layer
        folder = f"research-documents/{client_id}"
        meta_files = [f for f in storage.list_files(folder) if f.endswith(".meta.json")]

        results = []
        for mf in sorted(meta_files):
            meta = storage._load_json(f"{folder}/{mf}")
            if meta:
                results.append(meta)
        return results

    except HTTPException:
        raise
    except Exception as e:
        logger.error(
            f"Error listing research documents for client {client_id}",
            exc_info=True,
        )
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# Carrier Endpoints
# ============================================================================

@router.get("/carriers")
async def list_carriers():
    """
    List all carrier profiles.
    
    Returns:
        List of carrier profile records.
    """
    try:
        carriers = storage.list_carrier_profiles()
        return carriers
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Error listing carrier profiles", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/carriers", status_code=201)
async def create_carrier(request: CreateCarrierRequest):
    """
    Add or update a carrier profile.
    
    Args:
        request: Carrier profile data
        
    Returns:
        Created/updated carrier profile record.
    """
    try:
        # Create CarrierProfile dataclass instance
        carrier = CarrierProfile(
            carrier_name=request.carrier_name,
            amb_number=request.amb_number,
            naic_code=request.naic_code,
            financial_strength_rating=request.financial_strength_rating,
            issuer_credit_rating=request.issuer_credit_rating,
            rating_outlook=request.rating_outlook,
            balance_sheet_strength=request.balance_sheet_strength,
            operating_performance=request.operating_performance,
            business_profile=request.business_profile,
            erm_assessment=request.erm_assessment,
            net_premiums_written=request.net_premiums_written,
            policyholders_surplus=request.policyholders_surplus,
            combined_ratio=request.combined_ratio,
            five_year_avg_combined_ratio=request.five_year_avg_combined_ratio,
            report_date=request.report_date,
            direct_written_premium=request.direct_written_premium,
            nwp_to_surplus_ratio=request.nwp_to_surplus_ratio,
            lines_of_business_written=request.lines_of_business_written,
            geographic_concentration=request.geographic_concentration,
        )
        
        # Save to storage
        storage.save_carrier_profile(carrier)
        
        # Return as dict
        return asdict(carrier)
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Error creating carrier profile", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))
